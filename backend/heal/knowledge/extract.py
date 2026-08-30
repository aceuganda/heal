"""Getting text out of an uploaded file.

Kept separate from ingest so the failure modes are explicit. The one that
matters clinically is a scanned PDF: it has pages but no text layer, so naive
extraction returns an empty string and would ingest a source that answers
nothing while still being citable. That is refused loudly here.
"""
import io
from dataclasses import dataclass

from heal.logger import get_logger

logger = get_logger(__name__)

TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".text"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}
SUPPORTED = TEXT_SUFFIXES | PDF_SUFFIXES | DOCX_SUFFIXES


class ExtractionError(ValueError):
    """The file could not be turned into text worth indexing."""


@dataclass
class Extracted:
    text: str
    pages: int = 0
    kind: str = "text"


def extract(data: bytes, filename: str) -> Extracted:
    """Extract text from an uploaded file, by extension."""
    suffix = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if suffix in TEXT_SUFFIXES:
        return Extracted(text=_decode(data), kind="text")
    if suffix in PDF_SUFFIXES:
        return _extract_pdf(data)
    if suffix in DOCX_SUFFIXES:
        return _extract_docx(data)

    raise ExtractionError(
        f"Unsupported file type '{suffix or filename}'. "
        f"Supported: {', '.join(sorted(SUPPORTED))}"
    )


def _decode(data: bytes) -> str:
    """Decode uploaded bytes, preferring correctness over cleverness.

    UTF-16 is attempted only behind a byte-order mark. Tried speculatively it
    succeeds on almost any even-length input and returns mojibake, which is
    worse than a wrong-but-legible fallback: the text would be indexed and
    become citable.

    latin-1 cannot fail, so it terminates the chain.
    """
    if data[:2] in (b"\xff\xfe", b"\xfe\xff"):
        try:
            return data.decode("utf-16")
        except UnicodeDecodeError:
            pass
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError:
        logger.info("Upload is not UTF-8; falling back to latin-1")
        return data.decode("latin-1")


def _extract_pdf(data: bytes) -> Extracted:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise ExtractionError("PDF support requires pypdf") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:  # noqa: BLE001 -- any parse failure is the same to a user
        raise ExtractionError(f"Could not read the PDF: {exc}") from exc

    text = "\n\n".join(p.strip() for p in pages if p.strip())
    if not text.strip():
        # Almost always a scan. Ingesting it would create an approved source
        # that contains nothing and answers nothing, under a citation.
        raise ExtractionError(
            f"This PDF has {len(pages)} page(s) but no extractable text. It is "
            "most likely a scan. Run OCR on it first, or upload a text version "
            "-- indexing it as-is would create a citable source with no content."
        )
    return Extracted(text=text, pages=len(pages), kind="pdf")


def _extract_docx(data: bytes) -> Extracted:
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ExtractionError("DOCX support requires python-docx") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"Could not read the DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Tables carry the dosages in most treatment guidelines, so they are read
    # explicitly rather than left to the paragraph walk, which skips them.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n\n".join(parts)
    if not text.strip():
        raise ExtractionError("The document contains no extractable text")
    return Extracted(text=text, kind="docx")
